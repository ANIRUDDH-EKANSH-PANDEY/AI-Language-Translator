import streamlit as st
import requests
import os
from io import BytesIO
from docx import Document

# Set Page Configuration
st.set_page_config(page_title="AI Translator", layout="centered")

# Title
st.title("🌍 AI-Based Language Translator")

# Subheader Below Title
st.subheader("Enter Text or Upload a File for Translation 📄")

# Custom CSS for Styling
st.markdown("""
    <style>
        .stTextArea textarea {
            background-color: black !important;
            color: white !important;
            font-size: 16px !important;
            border-radius: 8px !important;
            padding: 10px !important;
        }
    </style>
""", unsafe_allow_html=True)

# Initialize Session State
if "input_text" not in st.session_state:
    st.session_state.input_text = ""
if "translated_text" not in st.session_state:
    st.session_state.translated_text = ""

# Callback function to clear text
def clear_text():
    st.session_state.input_text = ""
    st.session_state.translated_text = ""

# ✅ Text Area for Manual Input
text_input = st.text_area("Enter text to translate:", value=st.session_state.input_text, key="input_text")

# ✅ File Upload for Text Extraction
uploaded_file = st.file_uploader("Upload a file (TXT, DOCX)", type=["txt", "docx"])

# 📄 Extract Text from Uploaded File
extracted_text = ""
if uploaded_file:
    file_ext = uploaded_file.name.split(".")[-1].lower()  # Convert to lowercase for consistency

    try:
        if file_ext == "txt":
            extracted_text = uploaded_file.getvalue().decode("utf-8")

        elif file_ext == "docx":
            doc = Document(uploaded_file)
            extracted_text = "\n".join(para.text for para in doc.paragraphs) or "⚠ No text found in the document."

        else:
            st.error("❌ Unsupported file format. Please upload a TXT, or DOCX file.")

    except Exception as e:
        st.error(f"⚠ Error processing file: {str(e)}")
        extracted_text = ""

    # ✅ Show extracted text in a separate disabled text area
    if extracted_text and "⚠" not in extracted_text:
        st.text_area("Extracted Text:", value=extracted_text, height=200, disabled=True)

# 🌍 Language Selection Dropdown
languages = {
    "English": "en", "Hindi": "hi", "French": "fr", "German": "de",
    "Russian": "ru", "Spanish": "es", "Japanese": "ja", "Korean": "ko", 
    "Italian": "it", "Portuguese": "pt","Arabic": "ar", "Dutch": "nl", 
    "Greek": "el", "Turkish": "tr","Thai": "th", "Bengali": "bn", 
    "Urdu": "ur", "Tamil": "ta"
}
target_language = st.selectbox("Select Target Language:", list(languages.keys()))

# 🎯 Translate and Clear Buttons in Same Row
col1, col2 = st.columns([1, 1])  # Two equal columns

with col1:
    if st.button("Translate"):
        if text_input.strip() or extracted_text.strip():
            text_to_translate = text_input if text_input.strip() else extracted_text
            response = requests.post("https://ai-language-translator.onrender.com/translate",
                                     json={"text": text_to_translate, "target_language": languages[target_language]})
            if response.status_code == 200:
                st.session_state.translated_text = response.json().get("translated_text", "Translation failed")
            else:
                st.session_state.translated_text = "Error in translation."

with col2:
    st.button("Clear", on_click=clear_text)

# ✅ Show Translated Text
if st.session_state.translated_text:
    st.write("### Translated Text:")
    st.success(st.session_state.translated_text)

# 📝 Download Translated Text File
if st.session_state.translated_text:
    translated_text = st.session_state.translated_text  # Use translated text
    file_ext = uploaded_file.name.split(".")[-1] if uploaded_file else "txt"
    translated_filename = f"translated.{file_ext}"

    try:
        if file_ext == "txt":
            # Save translated text as plain text
            translated_bytes = translated_text.encode("utf-8")
            mime_type = "text/plain"

        elif file_ext == "docx":
            # Create a DOCX file from the translated text
            translated_doc = BytesIO()
            doc = Document()
            for line in translated_text.split("\n"):  # Preserve line breaks
                doc.add_paragraph(line)
            doc.save(translated_doc)
            translated_doc.seek(0)
            translated_bytes = translated_doc.read()
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        else:
            st.error("❌ Unsupported file format.")
            translated_bytes = None

        # Download button for the translated content
        if translated_bytes:
            st.download_button(
                label="📥 Download Translated Text",
                data=translated_bytes,
                file_name=translated_filename,
                mime=mime_type
            )

    except Exception as e:
        st.error(f"⚠ Something went wrong while generating the file: {str(e)}")
